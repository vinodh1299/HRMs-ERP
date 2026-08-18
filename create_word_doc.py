import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_architecture_doc(filename):
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles & Colors
    PRIMARY = RGBColor(0, 52, 112)       # Deep Navy
    SECONDARY = RGBColor(0, 163, 224)    # Cyan Blue
    TEXT_DARK = RGBColor(30, 41, 59)     # Slate Dark
    MUTED = RGBColor(100, 116, 139)      # Muted Slate

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Asian Christian Academy of India\nHRMS-ERP & Self-Hosted AI Engine Blueprint")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY
    
    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Self-Hosted Open-Weight Strategy, Infrastructure Engineering & Agent Portfolio\nVersion 1.0.0 · Document Review: August 2026")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = MUTED
    
    doc.add_paragraph() # Spacer

    # Helper function for section headings
    def add_heading(text, level=1):
        h = doc.add_paragraph()
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(15 if level==1 else 12.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(6)
        return h

    # Section 1: Executive Summary
    add_heading("1. Executive Summary & Self-Hosted AI Vision", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "This master architectural blueprint defines the technical specification, system design, agent portfolio, and infrastructure stack for building a 100% self-hosted, privately fine-tuned AI Engine for Asian Christian Academy of India. "
        "The system operates with total independence—eliminating third-party cloud API dependencies (zero OpenAI, zero Claude, zero Gemini API keys)—guaranteeing complete enterprise data privacy, zero per-token API costs, and low latency (30–150 ms TTFT) over a self-hosted open-weight model runtime."
    )
    r.font.name = 'Arial'
    r.font.size = Pt(10)

    # Section 2: Master Architecture Chart
    add_heading("2. Master System Architecture Chart (Self-Hosted Open Weights)", level=1)
    
    chart_box = doc.add_paragraph()
    chart_box.paragraph_format.space_before = Pt(6)
    chart_box.paragraph_format.space_after = Pt(6)
    
    chart_text = (
        "+-----------------------------------------------------------------------------------------+\n"
        "|  1. GLOBAL CLIENT UI ENGINE                                                             |\n"
        "|  Flutter 3.x (Dart) | Riverpod & GoRouter | Self-hosted STT + Platform TTS              |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                        |                                                  \n"
        "                                        v                                                  \n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|  2. ACA INDIA HRMS-ERP APPLICATION MODULES                                              |\n"
        "|  Dashboard (Presence/Polls) | Me (Attendance/Leaves) | Helpdesk (6 Depts) | Teams Chat |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                        |                                                  \n"
        "                                        v                                                  \n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|  3. EDGE & API GATEWAY                                                                  |\n"
        "|  nginx Edge (TLS 1.3, WAF, Rate Limit) -> Node.js Express API (JWT + RBAC Policy Check) |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                        |                                                  \n"
        "                                        v                                                  \n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|  4. AI AGENT ECOSYSTEM                                                                  |\n"
        "|  RAG Policy Agent | Action Agent | (Planned: Payroll | Attendance Audit | Vision OCR)    |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                        |                                                  \n"
        "                                        v                                                  \n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|  5. SELF-HOSTED INFERENCE & DATA                                                        |\n"
        "|  vLLM / llama.cpp Serving (LoRA Adapters) <-> Postgres 16 + pgvector (Unified DB)        |\n"
        "+-----------------------------------------------------------------------------------------+\n"
        "                                        |                                                  \n"
        "                                        v                                                  \n"
        "+-----------------------------------------------------------------------------------------+\n"
        "|  CROSS-CUTTING: SECURITY & OPERATIONS                                                   |\n"
        "|  SSO / IdP Auth | BullMQ + Redis Job Queue | Metrics & Tracing | Encrypted Backups   |\n"
        "+-----------------------------------------------------------------------------------------+"
    )
    
    r_chart = chart_box.add_run(chart_text)
    r_chart.font.name = 'Courier New'
    r_chart.font.size = Pt(8)
    r_chart.font.bold = True
    r_chart.font.color.rgb = PRIMARY

    # Section 3: Custom AI Development Stack & Codebase Engineering
    add_heading("3. Custom AI Infrastructure & Engineering Stack", level=1)
    
    tech_stack = [
        ("1. Model Fine-Tuning Framework (Python, PyTorch & LoRA)", "The core intelligence engine uses open-weight base models (e.g. Llama 3 / Gemma 2) fine-tuned on Asian Christian Academy of India HR policies, leave rules, and SOP documents using PyTorch, HuggingFace Transformers, and PEFT / LoRA (Low-Rank Adaptation)."),
        ("2. Dual-Tier Self-Hosted Inference Runtime (vLLM & llama.cpp)", "Model serving is structured in two operational tiers: vLLM (PagedAttention) provides primary GPU serving for high concurrent campus traffic (achieving 30–150 ms time-to-first-token), with llama.cpp (GGUF Quantization) acting as a low-power CPU fallback server."),
        ("3. Private Self-Hosted STT & Native Platform TTS", "To preserve complete data privacy, client voice input uses self-hosted whisper.cpp / faster-whisper speech-to-text models hosted on the local gateway, paired with native platform SpeechSynthesis voice-over engines on client devices."),
        ("4. Edge Security & TLS Termination (nginx & Express API Gateway)", "All traffic from mobile, web, and campus devices is encrypted in transit using TLS 1.3 HTTPS terminated at an nginx / Cloudflare edge node (with WAF and rate limiting) before routing to the Node.js / Express API service with JWT user authentication and RBAC policy enforcement."),
        ("5. Unified Relational & Vector Storage (Postgres 16 + pgvector)", "Relational ERP tables (employees, leaves, tickets, audit logs) and 1536-dimensional RAG document embeddings (knowledge_chunks) are stored in a unified Postgres 16 + pgvector database, guaranteeing transactional consistency and simplified single-node backup routines (RPO 24h / RTO 4h)."),
        ("6. Cross-Cutting Operations (BullMQ, Redis & Security Services)", "Long-running background tasks are managed via a BullMQ + Redis job queue. System health is monitored through open metrics and structured logging, backed by automated nightly AES-256 encrypted backups.")
    ]

    for title, desc in tech_stack:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_before = Pt(3); p_t.paragraph_format.space_after = Pt(3)
        r_title = p_t.add_run(f"• {title}: ")
        r_title.font.name = 'Arial'; r_title.font.size = Pt(9.5); r_title.font.bold = True; r_title.font.color.rgb = PRIMARY
        r_desc = p_t.add_run(desc)
        r_desc.font.name = 'Arial'; r_desc.font.size = Pt(9.5); r_desc.font.color.rgb = TEXT_DARK

    # Section 4: AI Agent Portfolio (Created vs Planned)
    add_heading("4. Asian Christian Academy AI Agent Ecosystem", level=1)
    
    active_agents = [
        ("📚 RAG Policy Search Agent (ragService.js)", "Performs grounded vector & BM25 search over indexed company policy documents."),
        ("⚡ Autonomous Action Agent (agentService.js)", "Converts user prompts into backend tool calls (createTicket, applyLeave, sendMessageToUserOrChannel)."),
        ("💬 Central Event Bus (ChatService)", "Cross-window synchronization service broadcasting live chat updates."),
        ("🎙️ Voice I/O Interface (VoiceHelper)", "Client-side voice interaction wrapper managing mic audio and TTS playback."),
        ("🛡️ HITL Safety Manager", "Staging service placing high-risk database mutations in agent_pending_actions for confirmation.")
    ]

    add_heading("4.1 Active Core Agents & Services", level=2)
    for title, desc in active_agents:
        p_a = doc.add_paragraph()
        r_t = p_a.add_run(f"• {title}: ")
        r_t.font.name = 'Arial'; r_t.font.size = Pt(9.5); r_t.font.bold = True; r_t.font.color.rgb = PRIMARY
        r_d = p_a.add_run(desc)
        r_d.font.name = 'Arial'; r_d.font.size = Pt(9.5); r_d.font.color.rgb = TEXT_DARK

    future_agents = [
        ("💰 Payroll & Tax Intelligence Agent", "Computes salary breakdowns, tax regime optimizations, and reimbursement claims."),
        ("⏰ Attendance & Shift Anomaly Agent", "Scans clock-in patterns and proactively suggests regularization workflows."),
        ("📸 Multi-Modal Vision & OCR Agent", "Extracts line items from expense receipts, medical bills, and onboarding ID cards."),
        ("🏷️ IT Diagnostic & Asset Specialist", "Tracks asset health, compares vendor pricing, and monitors PO approvals (> ₹1 Lakh)."),
        ("🔍 Hierarchical Escalation Agent", "Traverses organizational reporting trees for complex cross-department escalations.")
    ]

    add_heading("4.2 Planned Expansion Agents", level=2)
    for title, desc in future_agents:
        p_f = doc.add_paragraph()
        r_t = p_f.add_run(f"• {title}: ")
        r_t.font.name = 'Arial'; r_t.font.size = Pt(9.5); r_t.font.bold = True; r_t.font.color.rgb = PRIMARY
        r_d = p_f.add_run(desc)
        r_d.font.name = 'Arial'; r_d.font.size = Pt(9.5); r_d.font.color.rgb = TEXT_DARK

    # Section 5: Programming Language Evaluation
    add_heading("5. Programming Language Evaluation for Self-Hosted AI Infrastructure", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Language", "Evaluation & Suitability", "Strengths for Self-Hosted AI"]
    for i, title in enumerate(hdr_titles):
        hdr_cells[i].text = title
        shading = parse_xml(r'<w:shd {} w:fill="003470"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.name = 'Arial'; run.font.size = Pt(9.5); run.font.bold = True; run.font.color.rgb = RGBColor(255, 255, 255)

    lang_data = [
        ("Rust", "5/5 - Recommended Core", "Memory safety without GC pauses, zero-cost abstractions, native C/CUDA interop, high-speed tensor libraries (Candle / Burn), WebAssembly compilation."),
        ("Python", "4/5 - Essential Fine-Tuning", "Unrivaled AI ecosystem (PyTorch, JAX, HuggingFace, vLLM). Standard language for dataset preparation and LoRA adapter fine-tuning."),
        ("C++ / CUDA", "4/5 - Low-Level GPU Kernels", "Foundation of GGUF, llama.cpp, and TensorRT. Maximum low-level GPU hardware control."),
        ("Mojo", "3/5 - Emerging AI Tech", "Python superset compiling to native C speed with high-performance matrix math capabilities.")
    ]

    for lang, star, desc in lang_data:
        row_cells = table.add_row().cells
        row_cells[0].text = lang
        row_cells[1].text = star
        row_cells[2].text = desc
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Arial'; run.font.size = Pt(9); run.font.color.rgb = TEXT_DARK

    doc.save(filename)
    print(f"Successfully generated Word document: {filename}")

if __name__ == "__main__":
    create_architecture_doc("/Users/acamedia/VINODH/KEKA CLONE/PROJECT_OVERALL_ARCHITECTURE.docx")
    create_architecture_doc("/Users/acamedia/VINODH/KEKA CLONE/docs/PROJECT_OVERALL_ARCHITECTURE.docx")
